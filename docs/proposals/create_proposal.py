#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate microjet proposal document (.docx) for 凌策公司"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Constants ──
BLUE = RGBColor(0x1A, 0x3C, 0x6E)
LIGHT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT_NAME = "微軟正黑體"
OUTPUT_PATH = r"C:\Users\B00325\Desktop\公司AI比賽\lingce-company\docs\proposals\microjet_提案書.docx"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, size=12, bold=False, color=None, font_name=FONT_NAME):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.space_before = Pt(24)
        p.space_after = Pt(12)
    else:
        p.space_before = Pt(16)
        p.space_after = Pt(8)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=22, bold=True, color=BLUE)
    elif level == 2:
        set_run_font(run, size=16, bold=True, color=LIGHT_BLUE)
    else:
        set_run_font(run, size=13, bold=True, color=BLUE)
    if level == 1:
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="1A3C6E"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    return p


def add_body_text(doc, text, bold=False, size=11, color=DARK_GRAY, alignment=None, space_after=6):
    p = doc.add_paragraph()
    p.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, bold_prefix="", size=11):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="360"/>')
    pPr.append(ind)
    if bold_prefix:
        run_bold = p.add_run("\u25A0 " + bold_prefix)
        set_run_font(run_bold, size=size, bold=True, color=BLUE)
        run_text = p.add_run(text)
        set_run_font(run_text, size=size, color=DARK_GRAY)
    else:
        run = p.add_run("\u25A0 " + text)
        set_run_font(run, size=size, color=DARK_GRAY)
    return p


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color=WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A3C6E")

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=10, color=DARK_GRAY)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "EDF2F9")
            else:
                set_cell_shading(cell, "FFFFFF")

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    return table


def add_page_break(doc):
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    run = p.add_run()
    brk = OxmlElement('w:br')
    brk.set(qn('w:type'), 'page')
    run._element.append(brk)


# ══════════════════════════════════════════════════════════════
#  MAIN DOCUMENT CREATION
# ══════════════════════════════════════════════════════════════

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = FONT_NAME
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

# ══════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u51CC\u7B56\u516C\u53F8 \u00D7 microjet")
set_run_font(run, size=32, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(8)
run = p.add_run("\u5408\u4F5C\u63D0\u6848\u66F8")
set_run_font(run, size=28, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(16)
p.space_after = Pt(16)
run = p.add_run("\u2501" * 30)
set_run_font(run, size=12, color=LIGHT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(8)
run = p.add_run("AI \u9A45\u52D5\u667A\u6167\u88FD\u9020\u54C1\u8CEA\u76E3\u63A7\u7CFB\u7D71")
set_run_font(run, size=20, bold=True, color=LIGHT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(24)
run = p.add_run("AI-Driven Smart Manufacturing Quality Monitoring System")
set_run_font(run, size=12, color=RGBColor(0x66, 0x66, 0x66))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(40)
run = p.add_run("\u63D0\u6848\u65E5\u671F\uFF1A2026 \u5E74 4 \u6708 14 \u65E5")
set_run_font(run, size=12, color=DARK_GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u7248\u672C\uFF1AV1.0")
set_run_font(run, size=10, color=RGBColor(0x99, 0x99, 0x99))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(20)
run = p.add_run("\u6A5F\u5BC6\u6587\u4EF6 \u2500 \u50C5\u4F9B microjet \u5167\u90E8\u8A55\u4F30\u4F7F\u7528")
set_run_font(run, size=9, bold=True, color=RGBColor(0xCC, 0x33, 0x33))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "\u76EE\u9304", level=1)

toc_items = [
    ("\u4E00\u3001\u516C\u53F8\u7C21\u4ECB", "3"),
    ("\u4E8C\u3001\u9700\u6C42\u7406\u89E3", "4"),
    ("\u4E09\u3001\u89E3\u6C7A\u65B9\u6848", "5"),
    ("\u56DB\u3001\u6280\u8853\u67B6\u69CB", "7"),
    ("\u4E94\u3001\u5BE6\u65BD\u8A08\u756B", "8"),
    ("\u516D\u3001\u9810\u671F\u6548\u76CA", "9"),
    ("\u4E03\u3001\u70BA\u4EC0\u9EBC\u9078\u64C7\u51CC\u7B56", "10"),
    ("\u516B\u3001\u806F\u7D61\u8CC7\u8A0A", "11"),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    run = p.add_run(f"  {item}")
    set_run_font(run, size=12, color=DARK_GRAY)
    run2 = p.add_run(f"  {'.' * 50}  {page}")
    set_run_font(run2, size=12, color=RGBColor(0x99, 0x99, 0x99))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
#  SECTION 1: COMPANY INTRO
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "\u4E00\u3001\u516C\u53F8\u7C21\u4ECB", level=1)

add_body_text(doc, "\u95DC\u65BC\u51CC\u7B56\u516C\u53F8", bold=True, size=14, color=BLUE, space_after=8)

add_body_text(doc,
    "\u51CC\u7B56\u516C\u53F8\u662F\u53F0\u7063\u9996\u5BB6 100% AI \u9A45\u52D5\u7684\u8EDF\u9AD4\u516C\u53F8\uFF0C\u63A1\u7528 Palantir AIP \u71DF\u904B\u6A21\u5F0F\uFF0C"
    "\u4EE5\u300C\u4EBA\u985E\u76E3\u7BA1 + AI Agent \u5354\u4F5C\u300D\u7684\u5275\u65B0\u67B6\u69CB\u5B8C\u6210\u6240\u6709\u696D\u52D9\u6D41\u7A0B\u3002\u6211\u5011\u76F8\u4FE1 AI \u4E0D\u50C5\u662F\u5DE5\u5177\uFF0C"
    "\u66F4\u662F\u5718\u968A\u6838\u5FC3\u6210\u54E1\u3002",
    space_after=10
)

add_body_text(doc, "\u6838\u5FC3\u80FD\u529B", bold=True, size=13, color=BLUE, space_after=6)

agents_cap = [
    ("10 \u500B\u5C08\u696D AI Agent\uFF1A", "\u6DB5\u84CB\u7B56\u7565\u898F\u5283\u3001\u7CFB\u7D71\u958B\u767C\u3001\u54C1\u8CEA\u6E2C\u8A66\u3001\u5BA2\u6236\u670D\u52D9\u7B49\u5B8C\u6574\u696D\u52D9\u6D41\u7A0B"),
    ("AI \u539F\u751F\u67B6\u69CB\uFF1A", "\u5F9E\u516C\u53F8\u6210\u7ACB\u7B2C\u4E00\u5929\u5373\u4EE5 AI \u70BA\u6838\u5FC3\uFF0C\u975E\u50B3\u7D71\u4F01\u696D\u7684\u6578\u4F4D\u8F49\u578B"),
    ("\u5FEB\u901F\u4EA4\u4ED8\u80FD\u529B\uFF1A", "AI \u9A45\u52D5\u7684\u958B\u767C\u6D41\u7A0B\u53EF\u5C07\u5C08\u6848\u9031\u671F\u7E2E\u77ED 60% \u4EE5\u4E0A"),
    ("\u96E2\u7DDA AI \u90E8\u7F72\uFF1A", "\u652F\u63F4\u672C\u5730\u7AEF\u6A21\u578B\u90E8\u7F72\uFF0C\u78BA\u4FDD\u5BA2\u6236\u8CC7\u6599\u5B89\u5168\u4E0D\u5916\u6D29"),
]

for prefix, text in agents_cap:
    add_bullet(doc, text, bold_prefix=prefix)

add_body_text(doc, "", space_after=4)
add_body_text(doc, "AI Agent \u5718\u968A\u6982\u89BD", bold=True, size=13, color=BLUE, space_after=6)

agent_table_data = [
    ["\u7B56\u7565\u898F\u5283 Agent", "\u5E02\u5834\u5206\u6790\u3001\u5546\u696D\u7B56\u7565\u5236\u5B9A", "\u9AD8\u968E\u6C7A\u7B56\u652F\u63F4"],
    ["\u7CFB\u7D71\u67B6\u69CB Agent", "\u6280\u8853\u67B6\u69CB\u8A2D\u8A08\u3001\u6280\u8853\u9078\u578B", "\u7CFB\u7D71\u8A2D\u8A08"],
    ["\u5168\u7AEF\u958B\u767C Agent", "\u524D\u5F8C\u7AEF\u7A0B\u5F0F\u958B\u767C", "\u7522\u54C1\u958B\u767C"],
    ["AI/ML Agent", "\u6A5F\u5668\u5B78\u7FD2\u6A21\u578B\u958B\u767C\u8207\u8A13\u7DF4", "\u667A\u6167\u529F\u80FD"],
    ["\u54C1\u8CEA\u6E2C\u8A66 Agent", "\u81EA\u52D5\u5316\u6E2C\u8A66\u3001\u54C1\u8CEA\u628A\u95DC", "\u54C1\u8CEA\u4FDD\u8B49"],
    ["\u8CC7\u5B89\u9632\u8B77 Agent", "\u8CC7\u8A0A\u5B89\u5168\u8A55\u4F30\u8207\u9632\u8B77", "\u5B89\u5168\u5408\u898F"],
    ["DevOps Agent", "CI/CD\u3001\u90E8\u7F72\u7DAD\u904B", "\u7CFB\u7D71\u7DAD\u904B"],
    ["\u6578\u64DA\u5206\u6790 Agent", "\u8CC7\u6599\u5206\u6790\u8207\u8996\u89BA\u5316", "\u6578\u64DA\u6D1E\u5BDF"],
    ["\u5BA2\u6236\u670D\u52D9 Agent", "\u5BA2\u6236\u9700\u6C42\u5C0D\u63A5\u8207\u652F\u63F4", "\u5BA2\u6236\u95DC\u4FC2"],
    ["\u5C08\u6848\u7BA1\u7406 Agent", "\u5C08\u6848\u6392\u7A0B\u8207\u9032\u5EA6\u8FFD\u8E64", "\u5C08\u6848\u7BA1\u7406"],
]

create_table(doc,
    ["Agent \u540D\u7A31", "\u4E3B\u8981\u8077\u80FD", "\u8CA0\u8CAC\u9818\u57DF"],
    agent_table_data,
    col_widths=[5, 7, 4]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
#  SECTION 2: NEEDS ANALYSIS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "\u4E8C\u3001\u9700\u6C42\u7406\u89E3", level=1)

add_body_text(doc,
    "microjet \u4F5C\u70BA\u53F0\u7063\u7CBE\u5BC6\u5FAE\u578B\u5674\u5C04\u79D1\u6280\u7684\u9818\u5C0E\u5EE0\u5546\uFF0C\u5728\u9AD8\u7CBE\u5EA6\u88FD\u9020\u9818\u57DF\u64C1\u6709\u6DF1\u539A\u7684\u6280\u8853\u5BE6\u529B\u3002"
    "\u7136\u800C\uFF0C\u96A8\u8457\u5E02\u5834\u7AF6\u722D\u52A0\u5287\u8207\u5BA2\u6236\u54C1\u8CEA\u8981\u6C42\u63D0\u9AD8\uFF0C\u50B3\u7D71\u88FD\u9020\u6D41\u7A0B\u9762\u81E8\u4EE5\u4E0B\u95DC\u9375\u6311\u6230\uFF1A",
    space_after=12
)

challenges = [
    ("\u54C1\u8CEA\u6AA2\u6E2C\u6548\u7387\u74F6\u9838", [
        "\u4EBA\u5DE5\u76EE\u6AA2\u8017\u6642\u4E14\u6613\u53D7\u75B2\u52DE\u5F71\u97FF\uFF0C\u6F0F\u6AA2\u7387\u96A8\u5DE5\u6642\u589E\u52A0\u800C\u4E0A\u5347",
        "\u5FAE\u578B\u5674\u5634\u7B49\u7CBE\u5BC6\u96F6\u4EF6\u7684\u7F3A\u9677\u96E3\u4EE5\u8089\u773C\u8FA8\u8B58",
        "\u6AA2\u6E2C\u6A19\u6E96\u4F9D\u8CF4\u500B\u4EBA\u7D93\u9A57\uFF0C\u4E00\u81F4\u6027\u4E0D\u8DB3",
    ]),
    ("\u826F\u7387\u512A\u5316\u56F0\u96E3", [
        "\u7F3A\u4E4F\u5373\u6642\u7684\u826F\u7387\u6578\u64DA\u8FFD\u8E64\u6A5F\u5236",
        "\u826F\u7387\u6CE2\u52D5\u539F\u56E0\u96E3\u4EE5\u5FEB\u901F\u5B9A\u4F4D",
        "\u6279\u6B21\u9593\u54C1\u8CEA\u5DEE\u7570\u5206\u6790\u8017\u8CBB\u5927\u91CF\u4EBA\u529B",
    ]),
    ("\u7522\u7DDA\u6578\u64DA\u5206\u6790\u4E0D\u8DB3", [
        "\u88FD\u7A0B\u53C3\u6578\u8207\u54C1\u8CEA\u6578\u64DA\u672A\u80FD\u6709\u6548\u6574\u5408",
        "\u6B77\u53F2\u6578\u64DA\u5229\u7528\u7387\u4F4E\uFF0C\u96E3\u4EE5\u6316\u6398\u6F5B\u5728\u6539\u5584\u9EDE",
        "\u5831\u8868\u7522\u51FA\u4EF0\u8CF4\u4EBA\u5DE5\u5F59\u6574\uFF0C\u6642\u6548\u6027\u4E0D\u4F73",
    ]),
    ("\u9810\u6E2C\u6027\u7DAD\u8B77\u7F3A\u53E3", [
        "\u8A2D\u5099\u6545\u969C\u591A\u70BA\u4E8B\u5F8C\u8655\u7406\uFF0C\u9020\u6210\u975E\u9810\u671F\u505C\u6A5F",
        "\u7DAD\u8B77\u6392\u7A0B\u4F9D\u8CF4\u56FA\u5B9A\u9031\u671F\uFF0C\u672A\u80FD\u4F9D\u64DA\u5BE6\u969B\u72C0\u614B\u8ABF\u6574",
        "\u8A2D\u5099\u7570\u5E38\u7684\u65E9\u671F\u5FB5\u5146\u672A\u80FD\u88AB\u7CFB\u7D71\u6027\u6355\u6349",
    ]),
]

for title, items in challenges:
    add_heading_styled(doc, title, level=2)
    for item in items:
        add_bullet(doc, item)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
#  SECTION 3: SOLUTIONS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "\u4E09\u3001\u89E3\u6C7A\u65B9\u6848", level=1)

add_body_text(doc,
    "\u91DD\u5C0D\u4E0A\u8FF0\u6311\u6230\uFF0C\u51CC\u7B56\u516C\u53F8\u63D0\u51FA\u4EE5\u4E0B\u4E09\u5927\u89E3\u6C7A\u65B9\u6848\u6A21\u7D44\uFF0C\u53EF\u7368\u7ACB\u90E8\u7F72\u6216\u6574\u5408\u904B\u4F5C\uFF1A",
    space_after=12
)

# Solution 1
add_heading_styled(doc, "\u6A21\u7D44\u4E00\uFF1AAI \u54C1\u8CEA\u76E3\u63A7\u7CFB\u7D71", level=2)

add_body_text(doc, "\uFF081\uFF09AI \u8996\u89BA\u6AA2\u6E2C", bold=True, size=12, color=BLUE, space_after=4)
for item in [
    "\u5C0E\u5165\u6DF1\u5EA6\u5B78\u7FD2\u5F71\u50CF\u8FA8\u8B58\u6A21\u578B\uFF0C\u81EA\u52D5\u6AA2\u6E2C\u7522\u54C1\u5916\u89C0\u7F3A\u9677",
    "\u652F\u63F4\u5FAE\u7C73\u7D1A\u7CBE\u5EA6\u7684\u7626\u75B5\u5075\u6E2C\uFF08\u522E\u50B7\u3001\u6C23\u6CE1\u3001\u5C3A\u5BF8\u504F\u5DEE\u7B49\uFF09",
    "\u6A21\u578B\u53EF\u6839\u64DA microjet \u7522\u54C1\u7279\u6027\u9032\u884C\u5BA2\u88FD\u5316\u8A13\u7DF4",
    "\u6AA2\u6E2C\u901F\u5EA6\u53EF\u9054\u6BCF\u79D2 50+ \u4EF6\uFF0C\u70BA\u4EBA\u5DE5\u76EE\u6AA2\u7684 10 \u500D\u4EE5\u4E0A",
]:
    add_bullet(doc, item)

add_body_text(doc, "\uFF082\uFF09\u5373\u6642\u826F\u7387\u76E3\u63A7 Dashboard", bold=True, size=12, color=BLUE, space_after=4)
for item in [
    "\u7522\u7DDA\u826F\u7387\u5373\u6642\u8996\u89BA\u5316\u5100\u8868\u677F",
    "\u81EA\u52D5\u7D71\u8A08\u5404\u6279\u6B21\u3001\u5404\u7522\u7DDA\u3001\u5404\u6642\u6BB5\u7684\u826F\u7387\u6578\u64DA",
    "\u7570\u5E38\u826F\u7387\u81EA\u52D5\u8B66\u793A\u901A\u77E5\uFF08Email / LINE / \u5EE0\u5167\u5EE3\u64AD\uFF09",
    "\u652F\u63F4\u6B77\u53F2\u826F\u7387\u8DA8\u52E2\u5C0D\u6BD4\u8207\u6839\u56E0\u5206\u6790",
]:
    add_bullet(doc, item)

# Solution 2
add_heading_styled(doc, "\u6A21\u7D44\u4E8C\uFF1A\u9810\u6E2C\u6027\u7DAD\u8B77\u7CFB\u7D71", level=2)

add_body_text(doc, "\uFF081\uFF09\u8A2D\u5099\u7570\u5E38\u5075\u6E2C", bold=True, size=12, color=BLUE, space_after=4)
for item in [
    "\u5373\u6642\u76E3\u63A7\u8A2D\u5099\u632F\u52D5\u3001\u6EAB\u5EA6\u3001\u96FB\u6D41\u7B49\u95DC\u9375\u53C3\u6578",
    "AI \u7570\u5E38\u5075\u6E2C\u6F14\u7B97\u6CD5\u81EA\u52D5\u8B58\u5225\u8A2D\u5099\u9000\u5316\u8DA8\u52E2",
    "\u6545\u969C\u9810\u8B66\u63D0\u524D 48-72 \u5C0F\u6642\u767C\u51FA\uFF0C\u9810\u7559\u7DAD\u4FEE\u6E96\u5099\u6642\u9593",
]:
    add_bullet(doc, item)

add_body_text(doc, "\uFF082\uFF09\u7DAD\u8B77\u6392\u7A0B\u6700\u4F73\u5316", bold=True, size=12, color=BLUE, space_after=4)
for item in [
    "\u6839\u64DA\u8A2D\u5099\u5BE6\u969B\u72C0\u614B\u52D5\u614B\u8ABF\u6574\u7DAD\u8B77\u6392\u7A0B",
    "\u6574\u5408\u5099\u54C1\u5EAB\u5B58\u7BA1\u7406\uFF0C\u81EA\u52D5\u5EFA\u8B70\u5099\u54C1\u63A1\u8CFC",
    "\u7DAD\u8B77\u6210\u672C\u8207\u505C\u6A5F\u640D\u5931\u7684\u6700\u4F73\u5316\u5E73\u8861\u8A08\u7B97",
]:
    add_bullet(doc, item)

# Solution 3
add_heading_styled(doc, "\u6A21\u7D44\u4E09\uFF1A\u7522\u7DDA\u6578\u64DA\u5206\u6790\u5E73\u53F0", level=2)

add_body_text(doc, "\uFF081\uFF09\u81EA\u52D5\u5831\u8868\u751F\u6210", bold=True, size=12, color=BLUE, space_after=4)
for item in [
    "\u6BCF\u65E5 / \u6BCF\u9031 / \u6BCF\u6708\u54C1\u8CEA\u5831\u8868\u81EA\u52D5\u7522\u51FA",
    "\u652F\u63F4\u81EA\u8A02 KPI \u6307\u6A19\u8207\u5831\u8868\u683C\u5F0F",
    "\u5831\u8868\u81EA\u52D5\u6D3E\u767C\u81F3\u76F8\u95DC\u90E8\u9580\u4E3B\u7BA1",
]:
    add_bullet(doc, item)

add_body_text(doc, "\uFF082\uFF09\u8DA8\u52E2\u5206\u6790\u8207\u6D1E\u5BDF", bold=True, size=12, color=BLUE, space_after=4)
for item in [
    "\u88FD\u7A0B\u53C3\u6578\u8207\u54C1\u8CEA\u7684\u95DC\u806F\u6027\u5206\u6790",
    "\u9577\u671F\u8DA8\u52E2\u9810\u6E2C\u8207\u5B63\u7BC0\u6027\u6CE2\u52D5\u8B58\u5225",
    "AI \u9A45\u52D5\u7684\u6539\u5584\u5EFA\u8B70\u81EA\u52D5\u751F\u6210",
]:
    add_bullet(doc, item)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
#  SECTION 4: TECHNICAL ARCHITECTURE
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "\u56DB\u3001\u6280\u8853\u67B6\u69CB", level=1)

add_body_text(doc,
    "\u672C\u7CFB\u7D71\u63A1\u7528\u6DF7\u5408\u5F0F\u67B6\u69CB\uFF0C\u7D50\u5408\u96F2\u7AEF AI \u80FD\u529B\u8207\u672C\u5730\u7AEF\u90E8\u7F72\uFF0C\u5728\u78BA\u4FDD\u8CC7\u6599\u5B89\u5168\u7684\u524D\u63D0\u4E0B"
    "\u63D0\u4F9B\u6700\u4F73\u7684 AI \u63A8\u8AD6\u6548\u80FD\u3002",
    space_after=12
)

add_heading_styled(doc, "\u6838\u5FC3\u6280\u8853\u5143\u4EF6", level=2)

tech_table = [
    ["Claude API", "\u81EA\u7136\u8A9E\u8A00\u8655\u7406\u8207\u5206\u6790\u5831\u544A\u751F\u6210", "\u96F2\u7AEF\uFF08\u52A0\u5BC6\u50B3\u8F38\uFF09"],
    ["MCP \u5354\u8B70", "AI Agent \u8207\u4F01\u696D\u7CFB\u7D71\u7684\u6A19\u6E96\u5316\u9023\u63A5\u5C64", "\u672C\u5730 + \u96F2\u7AEF"],
    ["\u672C\u5730 AI \u6A21\u578B", "\u8996\u89BA\u6AA2\u6E2C\u3001\u7570\u5E38\u5075\u6E2C\u7B49\u63A8\u8AD6\u4EFB\u52D9", "\u5B8C\u5168\u96E2\u7DDA\u90E8\u7F72"],
    ["RAG \u77E5\u8B58\u5EAB", "\u88FD\u7A0B\u77E5\u8B58\u5EAB\u3001\u7DAD\u4FEE\u624B\u518A\u3001\u6B77\u53F2\u6848\u4F8B\u6AA2\u7D22", "\u672C\u5730\u90E8\u7F72"],
    ["\u5373\u6642\u4E32\u6D41\u5F15\u64CE", "\u8A2D\u5099\u611F\u6E2C\u5668\u6578\u64DA\u5373\u6642\u6536\u96C6\u8207\u8655\u7406", "\u672C\u5730\u90E8\u7F72"],
    ["Dashboard \u524D\u7AEF", "React + \u8CC7\u6599\u8996\u89BA\u5316", "\u672C\u5730\u5167\u7DB2"],
]

create_table(doc,
    ["\u6280\u8853\u5143\u4EF6", "\u529F\u80FD\u8AAA\u660E", "\u90E8\u7F72\u4F4D\u7F6E"],
    tech_table,
    col_widths=[4, 7, 4]
)

add_body_text(doc, "", space_after=8)

add_heading_styled(doc, "\u8CC7\u6599\u5B89\u5168\u6A5F\u5236", level=2)

for item in [
    "\u5DE5\u5EE0\u751F\u7522\u6578\u64DA\u5168\u7A0B\u7559\u5728\u672C\u5730\u7AEF\uFF0C\u4E0D\u4E0A\u50B3\u96F2\u7AEF",
    "\u8996\u89BA\u6AA2\u6E2C AI \u6A21\u578B\u65BC\u672C\u5730 GPU \u4F3A\u670D\u5668\u96E2\u7DDA\u63A8\u8AD6",
    "\u50C5\u975E\u654F\u611F\u7684\u7D71\u8A08\u6458\u8981\uFF08\u7D93\u5BA2\u6236\u6388\u6B0A\uFF09\u53EF\u900F\u904E\u52A0\u5BC6\u901A\u9053\u50B3\u8F38\u81F3\u96F2\u7AEF\u9032\u884C\u9032\u968E\u5206\u6790",
    "\u652F\u63F4\u6C23\u96D9\u7DB2\u8DEF\uFF08Air-gapped Network\uFF09\u90E8\u7F72\u6A21\u5F0F",
    "\u6240\u6709\u7CFB\u7D71\u64CD\u4F5C\u7559\u6709\u5B8C\u6574\u7A3D\u6838\u65E5\u8A8C",
]:
    add_bullet(doc, item)

add_body_text(doc, "", space_after=4)

add_heading_styled(doc, "\u7CFB\u7D71\u67B6\u69CB\u793A\u610F", level=2)

arch_table = [
    ["\u611F\u6E2C\u5668\u5C64", "\u6EAB\u5EA6 / \u632F\u52D5 / \u96FB\u6D41 / \u5F71\u50CF\u611F\u6E2C\u5668", "\u6578\u64DA\u63A1\u96C6"],
    ["\u908A\u7DE3\u904B\u7B97\u5C64", "\u672C\u5730 AI \u63A8\u8AD6\u4F3A\u670D\u5668 (GPU)", "\u5373\u6642\u5206\u6790"],
    ["\u61C9\u7528\u670D\u52D9\u5C64", "\u54C1\u8CEA\u76E3\u63A7 / \u9810\u6E2C\u7DAD\u8B77 / \u5831\u8868\u5F15\u64CE", "\u696D\u52D9\u908F\u8F2F"],
    ["\u5C55\u793A\u5C64", "Web Dashboard / \u884C\u52D5\u7AEF / \u8B66\u793A\u7CFB\u7D71", "\u4EBA\u6A5F\u4ECB\u9762"],
    ["\u77E5\u8B58\u5C64", "RAG \u77E5\u8B58\u5EAB / \u6B77\u53F2\u6848\u4F8B\u5EAB", "\u667A\u6167\u6C7A\u7B56"],
]

create_table(doc,
    ["\u67B6\u69CB\u5C64\u7D1A", "\u5143\u4EF6", "\u529F\u80FD"],
    arch_table,
    col_widths=[4, 7, 4]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
#  SECTION 5: IMPLEMENTATION PLAN
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "\u4E94\u3001\u5BE6\u65BD\u8A08\u756B", level=1)

add_body_text(doc,
    "\u51CC\u7B56\u516C\u53F8\u4EE5 AI \u9A45\u52D5\u7684\u654F\u6377\u958B\u767C\u6A21\u5F0F\uFF0C\u53EF\u5728 4 \u9031\u5167\u5B8C\u6210\u7CFB\u7D71\u5F9E\u9700\u6C42\u5230\u4E0A\u7DDA\u7684\u5B8C\u6574\u6D41\u7A0B\u3002",
    space_after=12
)

phase_table = [
    ["Phase 1", "\u9700\u6C42\u78BA\u8A8D\u8207\u6578\u64DA\u76E4\u9EDE", "1 \u9031", "\u7B2C 1 \u9031"],
    ["Phase 2", "\u6A21\u578B\u8A13\u7DF4\u8207\u7CFB\u7D71\u958B\u767C", "2 \u9031", "\u7B2C 2-3 \u9031"],
    ["Phase 3", "\u8A66\u7522\u7DDA\u90E8\u7F72\u8207\u512A\u5316", "1 \u9031", "\u7B2C 4 \u9031"],
]

create_table(doc,
    ["\u968E\u6BB5", "\u5DE5\u4F5C\u5167\u5BB9", "\u6642\u7A0B", "\u6392\u7A0B"],
    phase_table,
    col_widths=[3, 6, 3, 3]
)

add_body_text(doc, "", space_after=12)

add_heading_styled(doc, "Phase 1\uFF1A\u9700\u6C42\u78BA\u8A8D\u8207\u6578\u64DA\u76E4\u9EDE\uFF08\u7B2C 1 \u9031\uFF09", level=2)
for item in [
    "\u8207 microjet \u6280\u8853\u5718\u968A\u9032\u884C\u9700\u6C42\u8A2A\u8AC7",
    "\u76E4\u9EDE\u73FE\u6709\u8A2D\u5099\u611F\u6E2C\u5668\u8207\u6578\u64DA\u4ECB\u9762",
    "\u8A55\u4F30\u73FE\u5834\u7DB2\u8DEF\u74B0\u5883\u8207\u786C\u9AD4\u9700\u6C42",
    "\u78BA\u8A8D AI \u6A21\u578B\u8A13\u7DF4\u6240\u9700\u7684\u6A19\u8A3B\u6578\u64DA",
    "\u5236\u5B9A\u8A73\u7D30\u7684\u7CFB\u7D71\u898F\u683C\u66F8",
]:
    add_bullet(doc, item)

add_heading_styled(doc, "Phase 2\uFF1A\u6A21\u578B\u8A13\u7DF4\u8207\u7CFB\u7D71\u958B\u767C\uFF08\u7B2C 2-3 \u9031\uFF09", level=2)
for item in [
    "\u6839\u64DA microjet \u7522\u54C1\u6A23\u672C\u8A13\u7DF4\u8996\u89BA\u6AA2\u6E2C\u6A21\u578B",
    "\u958B\u767C\u8A2D\u5099\u7570\u5E38\u5075\u6E2C\u6F14\u7B97\u6CD5",
    "\u5EFA\u69CB\u5373\u6642\u76E3\u63A7 Dashboard",
    "\u6574\u5408 RAG \u77E5\u8B58\u5EAB\u8207\u5831\u8868\u5F15\u64CE",
    "\u9032\u884C\u7CFB\u7D71\u6574\u5408\u6E2C\u8A66",
]:
    add_bullet(doc, item)

add_heading_styled(doc, "Phase 3\uFF1A\u8A66\u7522\u7DDA\u90E8\u7F72\u8207\u512A\u5316\uFF08\u7B2C 4 \u9031\uFF09", level=2)
for item in [
    "\u65BC\u6307\u5B9A\u8A66\u7522\u7DDA\u5B89\u88DD\u8207\u90E8\u7F72\u7CFB\u7D71",
    "\u73FE\u5834\u74B0\u5883\u53C3\u6578\u6821\u6E96\u8207\u8ABF\u6574",
    "\u64CD\u4F5C\u4EBA\u54E1\u6559\u80B2\u8A13\u7DF4",
    "\u7CFB\u7D71\u6548\u80FD\u9A57\u8B49\u8207\u512A\u5316",
    "\u4EA4\u4ED8\u7CFB\u7D71\u64CD\u4F5C\u624B\u518A\u8207\u7DAD\u8B77\u6587\u4EF6",
]:
    add_bullet(doc, item)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
#  SECTION 6: EXPECTED BENEFITS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "\u516D\u3001\u9810\u671F\u6548\u76CA", level=1)

add_body_text(doc,
    "\u6839\u64DA\u51CC\u7B56\u516C\u53F8\u904E\u5F80\u5C08\u6848\u7D93\u9A57\u8207\u696D\u754C\u6A19\u7AFF\u6578\u64DA\uFF0C\u5C0E\u5165\u672C\u7CFB\u7D71\u5F8C\u9810\u671F\u53EF\u5E36\u4F86\u4EE5\u4E0B\u6548\u76CA\uFF1A",
    space_after=12
)

benefit_table = [
    ["\u7522\u54C1\u826F\u7387", "\u63D0\u5347 15% \u4EE5\u4E0A", "AI \u8996\u89BA\u6AA2\u6E2C\u7CBE\u6E96\u8FA8\u8B58\u7626\u75B5\uFF0C\u6E1B\u5C11\u6F0F\u6AA2\u8207\u8AA4\u5224"],
    ["\u6AA2\u6E2C\u901F\u5EA6", "\u63D0\u5347 10 \u500D", "\u81EA\u52D5\u5316\u6AA2\u6E2C\u53D6\u4EE3\u4EBA\u5DE5\u76EE\u6AA2\uFF0C\u5927\u5E45\u63D0\u5347\u7522\u80FD"],
    ["\u8A2D\u5099\u6545\u969C\u9810\u8B66", "\u6E96\u78BA\u7387 90%+", "\u63D0\u524D 48-72 \u5C0F\u6642\u9810\u8B66\uFF0C\u9810\u7559\u7DAD\u4FEE\u6E96\u5099\u6642\u9593"],
    ["\u7522\u7DDA\u505C\u6A5F\u6642\u9593", "\u964D\u4F4E 30%", "\u9810\u6E2C\u6027\u7DAD\u8B77\u6E1B\u5C11\u975E\u9810\u671F\u505C\u6A5F\u4E8B\u4EF6"],
    ["\u5831\u8868\u7522\u51FA\u6642\u9593", "\u6578\u5C0F\u6642 \u2192 \u5206\u9418", "\u81EA\u52D5\u5316\u5831\u8868\u751F\u6210\uFF0C\u5373\u6642\u6578\u64DA\u6D1E\u5BDF"],
    ["\u54C1\u8CEA\u4E00\u81F4\u6027", "\u986F\u8457\u63D0\u5347", "AI \u6AA2\u6E2C\u4E0D\u53D7\u75B2\u52DE\u5F71\u97FF\uFF0C\u6A19\u6E96\u59CB\u7D42\u4E00\u81F4"],
]

create_table(doc,
    ["\u6548\u76CA\u6307\u6A19", "\u9810\u671F\u6539\u5584", "\u8AAA\u660E"],
    benefit_table,
    col_widths=[4, 4, 8]
)

add_body_text(doc, "", space_after=12)

add_heading_styled(doc, "\u6295\u8CC7\u56DE\u5831\u4F30\u7B97", level=2)

add_body_text(doc,
    "\u4EE5\u4E2D\u578B\u7522\u7DDA\u898F\u6A21\u4F30\u7B97\uFF0C\u5C0E\u5165\u672C\u7CFB\u7D71\u5F8C\uFF1A",
    space_after=8
)

for item in [
    "\u826F\u7387\u63D0\u5347\u5E36\u4F86\u7684\u76F4\u63A5\u6210\u672C\u7BC0\u7701",
    "\u6E1B\u5C11\u4EBA\u5DE5\u6AA2\u6E2C\u4EBA\u529B\u6210\u672C",
    "\u964D\u4F4E\u8A2D\u5099\u975E\u9810\u671F\u6545\u969C\u7684\u640D\u5931",
    "\u63D0\u5347\u6574\u9AD4\u8A2D\u5099\u6548\u7387\uFF08OEE\uFF09",
]:
    add_bullet(doc, item)

add_body_text(doc,
    "\u9810\u671F\u6295\u8CC7\u56DE\u6536\u671F\u70BA 6-9 \u500B\u6708\u3002",
    bold=True, size=12, color=BLUE, space_after=12
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
#  SECTION 7: WHY LINGCE
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "\u4E03\u3001\u70BA\u4EC0\u9EBC\u9078\u64C7\u51CC\u7B56", level=1)

why_items = [
    ("AI \u539F\u751F\u67B6\u69CB",
     "\u51CC\u7B56\u4E26\u975E\u50B3\u7D71\u8EDF\u9AD4\u516C\u53F8\u8F49\u578B\u505A AI\uFF0C\u800C\u662F\u5F9E\u7B2C\u4E00\u5929\u5C31\u4EE5 AI \u70BA\u6838\u5FC3\u5EFA\u69CB\u7684\u4F01\u696D\u3002"
     "\u6211\u5011\u7684\u6BCF\u4E00\u500B\u6D41\u7A0B\u3001\u6BCF\u4E00\u884C\u7A0B\u5F0F\u78BC\u90FD\u878D\u5165\u4E86 AI \u601D\u7DAD\uFF0C\u80FD\u70BA\u5BA2\u6236\u6253\u9020\u771F\u6B63\u667A\u6167\u5316\u7684\u89E3\u6C7A\u65B9\u6848\u3002"),
    ("\u96E2\u7DDA\u7CFB\u7D71\u4FDD\u969C\u5DE5\u5EE0\u8CC7\u5B89",
     "\u6211\u5011\u6DF1\u77E5\u88FD\u9020\u696D\u5C0D\u8CC7\u6599\u5B89\u5168\u7684\u56B4\u683C\u8981\u6C42\u3002\u672C\u7CFB\u7D71\u652F\u63F4\u5B8C\u5168\u96E2\u7DDA\u90E8\u7F72\uFF0C"
     "\u6240\u6709\u751F\u7522\u6578\u64DA\u7559\u5728\u5DE5\u5EE0\u5167\u90E8\uFF0C\u4E0D\u9700\u9023\u63A5\u5916\u90E8\u7DB2\u8DEF\u5373\u53EF\u904B\u4F5C\uFF0C\u5FB9\u5E95\u6D88\u9664\u8CC7\u6599\u5916\u6D29\u98A8\u96AA\u3002"),
    ("\u5FEB\u901F\u90E8\u7F72\u80FD\u529B",
     "\u900F\u904E AI \u9A45\u52D5\u7684\u958B\u767C\u6D41\u7A0B\uFF0C\u6211\u5011\u53EF\u4EE5\u5728 4 \u9031\u5167\u5B8C\u6210\u5F9E\u9700\u6C42\u5230\u4E0A\u7DDA\u7684\u5B8C\u6574\u9031\u671F\u3002"
     "\u76F8\u8F03\u50B3\u7D71\u5C08\u6848\u52D5\u8F12 3-6 \u500B\u6708\u7684\u6642\u7A0B\uFF0C\u51CC\u7B56\u7684\u6548\u7387\u512A\u52E2\u986F\u8457\u3002"),
    ("Palantir AIP \u5354\u4F5C\u6A21\u5F0F",
     "\u51CC\u7B56\u63A1\u7528 Palantir \u7D1A\u5225\u7684 AI \u5354\u4F5C\u5E73\u53F0\uFF0C10 \u500B\u5C08\u696D AI Agent \u5354\u540C\u4F5C\u696D\uFF0C"
     "\u78BA\u4FDD\u6BCF\u500B\u74B0\u7BC0\u90FD\u6709\u5C08\u696D AI \u54C1\u8CEA\u628A\u95DC\uFF0C\u4EA4\u4ED8\u7269\u54C1\u8CEA\u7A69\u5B9A\u53EF\u9760\u3002"),
    ("\u6301\u7E8C\u512A\u5316\u627F\u8AFE",
     "\u7CFB\u7D71\u4E0A\u7DDA\u4E0D\u662F\u7D42\u9EDE\u3002\u6211\u5011\u7684 AI Agent \u6703\u6301\u7E8C\u5B78\u7FD2 microjet \u7684\u751F\u7522\u6578\u64DA\uFF0C"
     "\u96A8\u8457\u6578\u64DA\u7D2F\u7A4D\u4E0D\u65B7\u512A\u5316\u6A21\u578B\u7CBE\u5EA6\uFF0C\u8B93\u7CFB\u7D71\u8D8A\u7528\u8D8A\u6E96\u78BA\u3002"),
]

for title, desc in why_items:
    add_heading_styled(doc, title, level=2)
    add_body_text(doc, desc, space_after=10)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
#  SECTION 8: CONTACT INFO
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "\u516B\u3001\u806F\u7D61\u8CC7\u8A0A", level=1)

add_body_text(doc, "", space_after=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(20)
run = p.add_run("\u51CC\u7B56\u516C\u53F8 LingCe Company")
set_run_font(run, size=20, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(4)
run = p.add_run("100% AI-Driven Software Company")
set_run_font(run, size=12, color=RGBColor(0x66, 0x66, 0x66))

add_body_text(doc, "", space_after=16)

contact_table = [
    ["\u516C\u53F8\u540D\u7A31", "\u51CC\u7B56\u516C\u53F8\uFF08LingCe Company\uFF09"],
    ["\u71DF\u904B\u6A21\u5F0F", "100% AI \u9A45\u52D5\u8EDF\u9AD4\u958B\u767C"],
    ["\u6838\u5FC3\u6280\u8853", "AI Agent \u5354\u4F5C\u5E73\u53F0\uFF08Palantir AIP \u6A21\u5F0F\uFF09"],
    ["\u670D\u52D9\u7BC7\u7587", "\u667A\u6167\u88FD\u9020\u3001AI \u54C1\u8CEA\u76E3\u63A7\u3001\u9810\u6E2C\u6027\u7DAD\u8B77"],
    ["\u806F\u7D61\u7A97\u53E3", "\u696D\u52D9\u767C\u5C55\u90E8"],
]

table = doc.add_table(rows=len(contact_table), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, value) in enumerate(contact_table):
    cell_label = table.rows[i].cells[0]
    cell_value = table.rows[i].cells[1]

    cell_label.text = ""
    p = cell_label.paragraphs[0]
    run = p.add_run(label)
    set_run_font(run, size=11, bold=True, color=WHITE)
    set_cell_shading(cell_label, "1A3C6E")

    cell_value.text = ""
    p = cell_value.paragraphs[0]
    run = p.add_run(value)
    set_run_font(run, size=11, color=DARK_GRAY)
    if i % 2 == 0:
        set_cell_shading(cell_value, "EDF2F9")
    else:
        set_cell_shading(cell_value, "FFFFFF")

for row in table.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(10)

tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
    f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
    f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
    f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
    f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)

add_body_text(doc, "", space_after=20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(20)
run = p.add_run("\u611F\u8B1D microjet \u7684\u4FE1\u4EFB\u8207\u8A55\u4F30")
set_run_font(run, size=14, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(8)
run = p.add_run("\u51CC\u7B56\u516C\u53F8\u671F\u5F85\u8207\u8CB4\u516C\u53F8\u651C\u624B\uFF0C\u4EE5 AI \u6280\u8853\u958B\u5275\u667A\u6167\u88FD\u9020\u65B0\u7D00\u5143")
set_run_font(run, size=12, color=DARK_GRAY)

# ── Footer ──
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u51CC\u7B56\u516C\u53F8 LingCe Company \u2500 \u6A5F\u5BC6\u6587\u4EF6")
    set_run_font(run, size=8, color=RGBColor(0x99, 0x99, 0x99))

# ── Save ──
doc.save(OUTPUT_PATH)
print(f"Proposal saved to: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")
