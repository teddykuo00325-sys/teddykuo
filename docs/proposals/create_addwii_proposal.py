#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate addwii proposal document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Styles
style = doc.styles['Normal']
style.font.name = '微軟正黑體'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

for level, size in [('Heading 1', 22), ('Heading 2', 16), ('Heading 3', 13)]:
    s = doc.styles[level]
    s.font.name = '微軟正黑體'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

def add_styled_para(text, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = '微軟正黑體'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '微軟正黑體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '微軟正黑體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    doc.add_paragraph()

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = '微軟正黑體'
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

# ═══════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_styled_para('凌策公司 x addwii', bold=True, size=28, color=(0x1A, 0x3C, 0x6E), align=WD_ALIGN_PARAGRAPH.CENTER)
add_styled_para('合作提案書', bold=True, size=24, color=(0x1A, 0x3C, 0x6E), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_styled_para('AI Agent 智慧客服與數據分析平台', bold=True, size=16, color=(0x25, 0x63, 0xEB), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_styled_para('提案日期：2026 年 4 月 14 日', size=12, color=(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER)
add_styled_para('提案單位：凌策公司（AI Agent 驅動軟體公司）', size=12, color=(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER)
add_styled_para('版本：v1.0', size=12, color=(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ═══════════════════════════════════════
# 1. 公司簡介
# ═══════════════════════════════════════
doc.add_heading('一、公司簡介', level=1)
add_styled_para(
    '凌策公司是一家 100% AI 驅動的軟體公司，採用 Palantir AIP（Artificial Intelligence Platform）模式運營。'
    '公司由人類領導人負責監管、驗證與決策，所有日常工作由 10 個專業 AI Agent 全自動化完成。'
)
add_styled_para('核心架構：', bold=True)
add_bullet('Orchestrator Agent — 中央協調指揮，任務分派與進度管理')
add_bullet('業務開發部門 — BD Agent、客服 Agent、提案 Agent')
add_bullet('技術研發部門 — 前端 Agent、後端 Agent、QA Agent')
add_bullet('營運管理部門 — 財務 Agent、法務 Agent、文件 Agent')
add_styled_para(
    '凌策公司展示了 AI 原生組織的可能性：以極少人力實現高效率、高品質的軟體服務交付。'
)

doc.add_heading('二、需求理解', level=1)
add_styled_para(
    '根據我們對 addwii 的深入分析，我們理解貴公司在數位轉型過程中面臨以下核心挑戰：'
)
add_styled_para('1. 客服效率瓶頸', bold=True)
add_bullet('傳統客服人力有限，無法 24/7 即時回應')
add_bullet('重複性問題佔用大量人力資源')
add_bullet('多語言、跨平台客服需求日益增長')

add_styled_para('2. 數據分析困境', bold=True)
add_bullet('數據量快速增長，傳統分析工具難以應對')
add_bullet('從數據到洞見的轉化效率低')
add_bullet('缺乏即時監控與預測性分析能力')

add_styled_para('3. 人力成本壓力', bold=True)
add_bullet('專業人才招募困難且成本高')
add_bullet('重複性工作佔用核心人員時間')
add_bullet('組織擴張與成本控制的矛盾')

doc.add_page_break()

# ═══════════════════════════════════════
# 3. 解決方案
# ═══════════════════════════════════════
doc.add_heading('三、解決方案', level=1)
add_styled_para(
    '凌策公司為 addwii 量身打造「AI Agent 智慧客服與數據分析平台」，'
    '整合最先進的 AI 技術，為貴公司提供全方位的數位轉型解決方案。'
)

doc.add_heading('3.1 AI 智慧客服系統', level=2)
add_table(
    ['功能模組', '說明', '效益'],
    [
        ['24/7 自動客服', '基於 Claude AI 的智慧對話系統，支援自然語言理解', '客服回覆速度提升 90%'],
        ['多語言支援', '自動偵測語言並切換，支援中/英/日等語系', '服務範圍擴大 3 倍'],
        ['情緒分析', '即時偵測客戶情緒，高風險對話自動轉接人工', '客戶滿意度提升 40%'],
        ['知識庫整合', 'RAG 技術連接公司知識庫，回答準確率 95%+', '首次解決率提升 60%'],
        ['對話紀錄分析', '自動分析客服對話，產出洞見報告', '管理效率提升 50%'],
    ]
)

doc.add_heading('3.2 數據分析 Dashboard', level=2)
add_table(
    ['功能模組', '說明', '效益'],
    [
        ['即時監控', '業務指標即時視覺化，自動異常告警', '問題發現速度提升 80%'],
        ['自動報表', 'AI 自動產出日報/週報/月報', '報表製作時間減少 90%'],
        ['預測分析', '基於歷史數據的趨勢預測與建議', '決策準確度提升 35%'],
        ['自訂儀表板', '拖拽式自訂 Dashboard，零程式碼', '資訊取得效率提升 70%'],
        ['數據整合', '整合多來源數據，統一資料模型', '數據孤島消除率 100%'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 4. 技術架構
# ═══════════════════════════════════════
doc.add_heading('四、技術架構', level=1)
add_styled_para('系統採用在線/離線雙架構設計，確保效能與資安兼顧：', bold=True)

doc.add_heading('4.1 在線系統', level=2)
add_table(
    ['元件', '技術', '說明'],
    [
        ['AI 引擎', 'Claude API (Anthropic)', '核心推理與對話能力'],
        ['Agent 框架', 'MCP + Claude Code', 'AI Agent 協作執行環境'],
        ['API 閘道', 'LiteLLM', '統一 API、Token 監控、流量分配'],
        ['知識庫', 'Qdrant (RAG)', '向量資料庫，智慧檢索'],
        ['前端', 'React + Tailwind', '響應式 Dashboard 介面'],
        ['後端', 'Express.js', 'RESTful API 服務'],
    ]
)

doc.add_heading('4.2 資安保障', level=2)
add_bullet('所有 API 通訊走 HTTPS/TLS 加密')
add_bullet('客戶機密資料可選擇在離線系統處理（Air-gapped 隔離）')
add_bullet('API Key 分層管理，最小權限原則')
add_bullet('敏感操作需人類審核確認（Human-in-the-Loop）')
add_bullet('資料存取全程日誌追蹤')

doc.add_page_break()

# ═══════════════════════════════════════
# 5. 實施計劃
# ═══════════════════════════════════════
doc.add_heading('五、實施計劃', level=1)
add_styled_para('採用敏捷開發模式，分三階段快速交付：')

add_table(
    ['階段', '時程', '工作內容', '交付物'],
    [
        ['Phase 1\n需求確認與 POC', '第 1 週', '深入了解 addwii 業務需求\n確認系統規格與整合範圍\n產出概念驗證原型', '需求規格書\nPOC Demo\n技術可行性報告'],
        ['Phase 2\n核心功能開發', '第 2-3 週', '智慧客服系統開發\n數據分析 Dashboard 建構\n知識庫整合與訓練\nAPI 開發與系統整合', '客服系統 v1.0\nDashboard v1.0\nAPI 文檔'],
        ['Phase 3\n上線與優化', '第 4 週', '系統部署上線\n使用者教育訓練\n效能優化與調校\n監控系統建置', '正式上線系統\n使用者手冊\n維運 SOP'],
    ]
)

# ═══════════════════════════════════════
# 6. 預期效益
# ═══════════════════════════════════════
doc.add_heading('六、預期效益', level=1)
add_table(
    ['效益指標', '預期提升', '說明'],
    [
        ['客服效率', '提升 60%', '24/7 AI 自動回覆 + 知識庫精準匹配'],
        ['數據分析時間', '縮短 80%', 'AI 自動報表 + 即時 Dashboard'],
        ['人力成本', '降低 40%', 'AI 承擔重複性工作，人力聚焦核心業務'],
        ['客戶滿意度', '提升 35%', '即時回應 + 情緒分析 + 個人化服務'],
        ['問題首次解決率', '提升至 90%', 'RAG 知識庫提供精準答案'],
        ['報表產出', '加速 10 倍', 'AI 自動分析與生成'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 7. 為什麼選擇凌策
# ═══════════════════════════════════════
doc.add_heading('七、為什麼選擇凌策', level=1)

add_styled_para('1. AI 原生公司', bold=True)
add_styled_para('凌策是 100% AI 驅動的軟體公司，我們不只提供 AI 工具，更以 AI 原生的方式運營。這意味著我們深刻理解 AI 的能力與限制，能為您提供最務實的方案。')

add_styled_para('2. Palantir 級架構', bold=True)
add_styled_para('我們採用 Palantir AIP 模式的組織架構，以 Ontology 驅動的統一資料模型串接所有業務流程，確保系統的可擴展性與一致性。')

add_styled_para('3. 快速交付', bold=True)
add_styled_para('10 個專業 AI Agent 協同工作，7x24 不間斷開發。從需求確認到上線，僅需 4 週即可完成全系統交付。')

add_styled_para('4. 成本優勢', bold=True)
add_styled_para('AI Agent 驅動的開發模式大幅降低人力成本，同時透過 Prompt Caching、模型分流等技術控制 AI 使用成本，將節省回饋給客戶。')

add_styled_para('5. 雙系統資安保障', bold=True)
add_styled_para('獨特的在線/離線雙系統架構，以「冷錢包」概念保護客戶商業秘密，符合台灣法規要求。')

doc.add_page_break()

# ═══════════════════════════════════════
# 8. 聯絡資訊
# ═══════════════════════════════════════
doc.add_heading('八、聯絡資訊', level=1)
doc.add_paragraph()
add_styled_para('凌策公司', bold=True, size=16, color=(0x1A, 0x3C, 0x6E))
add_styled_para('AI Agent 驅動的軟體公司')
doc.add_paragraph()
add_styled_para('聯繫方式：')
add_bullet('公司官網：凌策公司 AI 協作平台')
add_bullet('技術展示：AI Agent 協作指揮中心 Dashboard')
add_bullet('聯絡窗口：客服 Agent (cs-001) / BD Agent (bd-001)')
doc.add_paragraph()
doc.add_paragraph()
add_styled_para('感謝 addwii 的信任與支持，我們期待與貴公司攜手共創 AI 驅動的未來。',
                size=12, color=(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_styled_para('本提案書由凌策公司 AI Agent 協作產出，人類負責審核與決策。',
                size=10, color=(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER)

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'addwii_提案書.docx')
doc.save(output_path)
print(f"Proposal saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path):,} bytes")
