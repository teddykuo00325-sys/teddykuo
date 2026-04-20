"""Generate professional proposal document for 維明."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BLUE = RGBColor(0x1A, 0x3C, 0x6E)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_BLUE_BG = "D6E4F0"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT_NAME = "微軟正黑體"
FONT_NAME_EN = "Calibri"

doc = Document()

# -- Default style setup --
style = doc.styles["Normal"]
font = style.font
font.name = FONT_NAME_EN
font.size = Pt(11)
font.color.rgb = DARK_GRAY
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

# Adjust margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, size=11, bold=False, color=DARK_GRAY, font_name=FONT_NAME):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = FONT_NAME_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_heading_styled(text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.space_before = Pt(24)
        p.space_after = Pt(12)
    else:
        p.space_before = Pt(16)
        p.space_after = Pt(8)
    run = p.add_run(text)
    size = {1: 20, 2: 16, 3: 13}.get(level, 14)
    set_run_font(run, size=size, bold=True, color=BLUE)
    # Add bottom border for level 1
    if level == 1:
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="1A3C6E"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    return p


def add_para(text, size=11, bold=False, color=DARK_GRAY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(text, level=0, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, size=11, bold=True, color=DARK_GRAY)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=False, color=DARK_GRAY)
    return p


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=11, bold=True, color=WHITE)
        set_cell_shading(cell, "1A3C6E")
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, size=10, bold=False, color=DARK_GRAY)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FA")
    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing
    return table


# ============================================================
# 1. COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("━" * 40)
set_run_font(run, size=14, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(8)
run = p.add_run("凌策公司 x 維明")
set_run_font(run, size=32, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(4)
run = p.add_run("合作提案書")
set_run_font(run, size=28, bold=True, color=BLUE)

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("━" * 40)
set_run_font(run, size=14, color=BLUE)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(12)
run = p.add_run("AI Agent 企業流程自動化解決方案")
set_run_font(run, size=18, bold=False, color=RGBColor(0x4A, 0x6F, 0xA5))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("提案日期：2026 年 4 月 15 日")
set_run_font(run, size=12, color=DARK_GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("文件編號：LC-PROP-2026-0415-WM")
set_run_font(run, size=10, color=RGBColor(0x88, 0x88, 0x88))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("機密文件 CONFIDENTIAL")
set_run_font(run, size=10, bold=True, color=RGBColor(0xCC, 0x33, 0x33))

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled("目 錄", level=1)

toc_items = [
    ("一、", "公司簡介"),
    ("二、", "需求理解"),
    ("三、", "解決方案"),
    ("四、", "技術架構"),
    ("五、", "實施計劃"),
    ("六、", "預期效益"),
    ("七、", "為什麼選擇凌策"),
    ("八、", "聯絡資訊"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.left_indent = Cm(2)
    run = p.add_run(f"{num}{title}")
    set_run_font(run, size=13, bold=False, color=BLUE)

doc.add_page_break()

# ============================================================
# 2. COMPANY INTRODUCTION
# ============================================================
add_heading_styled("一、公司簡介", level=1)

add_para(
    "凌策公司（LingCe Company）是一家 100% AI 驅動的軟體公司，採用 Palantir AIP（Artificial Intelligence Platform）"
    "模式，以「人類監管 + AI Agent 協作」的創新運營架構完成所有業務。我們相信，AI 不僅是工具，更是團隊的核心成員。"
)

add_heading_styled("核心理念", level=2)

add_bullet("AI 原生營運：", bold_prefix="")
add_para(
    "凌策公司自創立之初即以 AI 為核心建構所有業務流程。不同於傳統公司在既有流程上疊加 AI 工具，"
    "我們從零開始設計以 AI Agent 為主體的工作模式，確保每個環節都能發揮 AI 的最大效能。",
    size=10, space_after=8
)

add_heading_styled("AI Agent 團隊", level=2)

add_para(
    "凌策公司擁有 10 個專業 AI Agent，各司其職，涵蓋軟體開發、系統架構、品質保證、專案管理、"
    "商業分析、客戶服務等面向，形成完整的 AI 協作生態系統。"
)

agent_data = [
    ["Agent 類型", "專業領域", "核心能力"],
    ["開發 Agent", "軟體工程", "全端開發、API 設計、程式碼審查"],
    ["架構 Agent", "系統設計", "微服務架構、雲端部署、效能優化"],
    ["QA Agent", "品質保證", "自動化測試、安全檢測、效能測試"],
    ["PM Agent", "專案管理", "需求分析、進度追蹤、資源調配"],
    ["BA Agent", "商業分析", "市場研究、數據分析、競爭分析"],
    ["CS Agent", "客戶服務", "技術支援、問題診斷、知識庫管理"],
    ["DevOps Agent", "維運部署", "CI/CD、監控告警、自動擴容"],
    ["Security Agent", "資安防護", "弱點掃描、合規檢查、事件回應"],
    ["Data Agent", "數據工程", "ETL 處理、資料建模、報表生成"],
    ["Doc Agent", "文件管理", "技術文檔、使用手冊、知識庫維護"],
]
make_table(agent_data[0], agent_data[1:], col_widths=[4, 4, 8])

doc.add_page_break()

# ============================================================
# 3. REQUIREMENT UNDERSTANDING
# ============================================================
add_heading_styled("二、需求理解", level=1)

add_para(
    "經過初步評估與了解，我們深刻理解維明作為台灣領先的企業服務與系統整合公司，"
    "在數位轉型過程中所面臨的核心挑戰。以下為我們對貴公司現況的分析與理解："
)

add_heading_styled("企業現況挑戰", level=2)

challenges = [
    ("業務流程人工重複操作", "大量日常業務流程仍依賴人工手動執行，包括資料登錄、報表彙整、郵件通知等重複性工作，"
     "不僅耗費人力資源，也容易因人為疏失產生錯誤。"),
    ("跨系統資料整合困難", "企業內部多套系統（ERP、CRM、HR 等）各自獨立運作，資料格式不一、缺乏統一接口，"
     "導致資訊孤島、決策資訊不即時。"),
    ("文件處理耗時", "合約審核、報價單生成、技術文件整理等文件作業，需要大量人力投入核對與格式化，"
     "處理週期過長影響業務效率。"),
    ("客服回覆效率低", "客戶諮詢量大，人工客服回覆速度有限，難以做到 24/7 即時回應，"
     "影響客戶滿意度與業務成交率。"),
]
for title, desc in challenges:
    add_bullet(title, bold_prefix="")
    add_para(desc, size=10, space_after=8)

add_heading_styled("轉型目標", level=2)
add_para(
    "維明期望透過 AI 技術的導入，實現業務流程的智慧化與自動化，降低營運成本的同時提升服務品質，"
    "並在台灣企業服務市場中建立差異化的競爭優勢。"
)

doc.add_page_break()

# ============================================================
# 4. SOLUTION
# ============================================================
add_heading_styled("三、解決方案", level=1)

add_para(
    "針對維明的需求，凌策公司提出以下三大核心解決方案模組，以 AI Agent 技術為基礎，"
    "打造全方位的企業流程自動化平台。"
)

# Solution 1
add_heading_styled("方案一：AI Agent 流程自動化", level=2)

add_para(
    "結合 RPA（機器人流程自動化）與 AI 智慧判斷能力，打造具備理解力與決策力的流程自動化系統。"
)

sol1_items = [
    ("RPA + AI 智慧結合：", "傳統 RPA 僅能處理固定規則流程，我們的 AI Agent 能理解語義、判斷例外狀況，"
     "自動處理 90% 以上的標準業務流程，並將異常案例智慧分流給人工處理。"),
    ("自動文件處理：", "AI Agent 自動識別、擷取、分類文件內容，支援合約、報價單、技術文件等多種格式，"
     "處理速度較人工提升 10 倍以上。"),
    ("智慧審批流程：", "根據企業規則自動進行初審、風險評估與合規檢查，僅需人工確認最終決策，"
     "大幅縮短審批週期。"),
]
for prefix, desc in sol1_items:
    add_bullet(desc, bold_prefix=prefix)

# Solution 2
add_heading_styled("方案二：跨系統數據整合", level=2)

add_para(
    "建構統一的數據整合層，打通企業內部各系統的資料壁壘。"
)

sol2_items = [
    ("API Gateway 統一接口：", "建立企業級 API Gateway，提供標準化的數據存取接口，"
     "支援 REST、GraphQL 等協定，無縫串接 ERP、CRM、HR 等系統。"),
    ("RAG 知識庫：", "運用檢索增強生成（Retrieval-Augmented Generation）技術，"
     "將企業文件、SOP、歷史案例建構為智慧知識庫，AI Agent 可即時查詢並生成精確回覆。"),
    ("即時數據同步：", "事件驅動架構確保各系統資料即時同步，消除資訊延遲，"
     "提供管理層即時的營運儀表板與決策支援。"),
]
for prefix, desc in sol2_items:
    add_bullet(desc, bold_prefix=prefix)

# Solution 3
add_heading_styled("方案三：AI 客服助手", level=2)

add_para(
    "部署 AI 驅動的智慧客服系統，提供 24/7 全天候的客戶服務體驗。"
)

sol3_items = [
    ("多渠道整合：", "支援網頁、LINE、Email 等多渠道客服接入，統一管理所有客戶互動。"),
    ("智慧路由：", "AI 自動判斷問題類型與緊急程度，精準分配至對應的處理流程或專業人員。"),
    ("持續學習：", "AI Agent 從每次互動中學習優化，回覆品質與準確度隨時間持續提升。"),
]
for prefix, desc in sol3_items:
    add_bullet(desc, bold_prefix=prefix)

doc.add_page_break()

# ============================================================
# 5. TECHNICAL ARCHITECTURE
# ============================================================
add_heading_styled("四、技術架構", level=1)

add_para(
    "凌策公司的解決方案建構於成熟且安全的技術架構之上，確保企業級的穩定性、安全性與擴展性。"
)

add_heading_styled("核心技術堆疊", level=2)

tech_data = [
    ["技術層", "技術方案", "說明"],
    ["AI 引擎", "Claude API", "Anthropic 頂級大語言模型，具備強大的理解與生成能力"],
    ["工具協定", "MCP (Model Context Protocol)", "標準化的 AI 工具調用協定，支援多元工具整合"],
    ["Agent 編排", "Agent Orchestration", "多 Agent 協作調度引擎，實現複雜任務的自動分解與執行"],
    ["資料安全", "在線/離線雙系統", "敏感資料離線處理，一般資料在線處理，確保企業資料安全"],
    ["部署方式", "混合雲架構", "支援公有雲與私有雲混合部署，符合企業合規要求"],
    ["監控管理", "即時監控儀表板", "全方位監控 Agent 運作狀態、效能指標與異常告警"],
]
make_table(tech_data[0], tech_data[1:], col_widths=[3.5, 5, 7.5])

add_heading_styled("資料安全架構", level=2)

add_para(
    "企業資料安全是我們的最高優先級。凌策公司採用「在線/離線雙系統」架構，確保資料處理的安全性："
)

security_items = [
    ("離線處理系統：", "涉及客戶個資、財務數據、營業秘密等敏感資料，全程在企業內部私有環境處理，"
     "資料不離開企業邊界。"),
    ("在線處理系統：", "一般性的文件生成、知識查詢、流程自動化等任務，透過加密連線使用雲端 AI 服務，"
     "兼顧效能與成本效益。"),
    ("存取控制：", "基於角色的存取控制（RBAC），完整的操作日誌記錄，"
     "確保每一筆資料存取皆可追溯與稽核。"),
]
for prefix, desc in security_items:
    add_bullet(desc, bold_prefix=prefix)

doc.add_page_break()

# ============================================================
# 6. IMPLEMENTATION PLAN
# ============================================================
add_heading_styled("五、實施計劃", level=1)

add_para(
    "凌策公司規劃為期 4 週的敏捷式導入計劃，採用漸進式部署策略，確保每個階段都有明確的交付物與驗收標準。"
)

plan_data = [
    ["階段", "時程", "主要工作項目", "交付物"],
    ["Phase 1\n流程盤點與 POC", "第 1 週", "- 現有流程深度訪談\n- 關鍵流程優先級排序\n- POC 環境建置\n- 概念驗證展示",
     "- 流程盤點報告\n- 優先級矩陣\n- POC 成果展示"],
    ["Phase 2\n核心 Agent 開發\n與系統整合", "第 2-3 週", "- 核心 AI Agent 開發\n- API Gateway 建置\n- RAG 知識庫構建\n- 系統整合測試",
     "- 核心 Agent 模組\n- 整合測試報告\n- 系統架構文件"],
    ["Phase 3\n全面部署與\n人員訓練", "第 4 週", "- 正式環境部署\n- 使用者教育訓練\n- 上線支援與調優\n- 維運文件交付",
     "- 正式上線系統\n- 訓練教材\n- 維運手冊\n- 驗收報告"],
]
make_table(plan_data[0], plan_data[1:], col_widths=[3.5, 2.5, 5.5, 4.5])

add_heading_styled("各階段詳細說明", level=2)

add_heading_styled("Phase 1：流程盤點與 POC（第 1 週）", level=3)
add_para(
    "深入了解維明現有業務流程，識別最具自動化價值的關鍵流程，並透過概念驗證展示 AI Agent 的實際效果，"
    "讓團隊在第一時間看到具體成果，建立導入信心。"
)

add_heading_styled("Phase 2：核心 Agent 開發與系統整合（第 2-3 週）", level=3)
add_para(
    "根據 POC 結果與優先級排序，開發核心 AI Agent 模組，建置 API Gateway 與 RAG 知識庫，"
    "並進行全面的系統整合測試，確保各模組穩定運作。"
)

add_heading_styled("Phase 3：全面部署與人員訓練（第 4 週）", level=3)
add_para(
    "將系統部署至正式環境，進行使用者教育訓練，確保維明團隊能獨立操作與管理 AI Agent 系統。"
    "同時提供完整的維運文件與持續支援計劃。"
)

doc.add_page_break()

# ============================================================
# 7. EXPECTED BENEFITS
# ============================================================
add_heading_styled("六、預期效益", level=1)

add_para(
    "根據凌策公司過往專案經驗與產業數據分析，導入 AI Agent 企業流程自動化解決方案後，"
    "維明可期待以下顯著效益："
)

benefit_data = [
    ["效益指標", "現況估計", "導入後預期", "提升幅度"],
    ["流程處理效率", "每案 2-4 小時", "每案 30-60 分鐘", "提升 70%"],
    ["文件處理時間", "每份 30-60 分鐘", "每份 5-10 分鐘", "縮短 85%"],
    ["人力成本", "基準值", "減少 50%", "降低 50%"],
    ["作業錯誤率", "平均 5-8%", "低於 1%", "降低 90%"],
    ["客服回覆時間", "平均 2-4 小時", "即時回覆（< 1 分鐘）", "提升 95%"],
    ["服務可用時間", "週一至週五\n09:00-18:00", "24/7 全天候", "覆蓋率 100%"],
]
make_table(benefit_data[0], benefit_data[1:], col_widths=[3.5, 3.5, 4.5, 3])

add_heading_styled("投資報酬分析", level=2)
add_para(
    "根據初步估算，導入 AI Agent 解決方案的投資，預計可在 3-6 個月內達到投資回收（ROI），"
    "長期而言將為維明帶來持續且可觀的成本節省與效率提升。"
)

add_heading_styled("無形效益", level=2)
intangible = [
    ("員工滿意度提升：", "將員工從重複性工作中釋放，投入更具價值的策略性任務。"),
    ("決策品質提升：", "即時、準確的數據分析支援管理層做出更好的商業決策。"),
    ("客戶體驗升級：", "更快速、更精準的服務回覆，提升客戶忠誠度與品牌形象。"),
    ("競爭優勢建立：", "率先導入 AI 自動化的企業，將在市場中取得顯著的先行者優勢。"),
]
for prefix, desc in intangible:
    add_bullet(desc, bold_prefix=prefix)

doc.add_page_break()

# ============================================================
# 8. WHY LINGCE
# ============================================================
add_heading_styled("七、為什麼選擇凌策", level=1)

add_para(
    "在眾多 AI 解決方案供應商中，凌策公司具備獨一無二的優勢，是維明推動 AI 轉型的最佳合作夥伴。"
)

add_heading_styled("AI 原生公司", level=2)
add_para(
    "凌策公司不是「導入 AI 的傳統公司」，而是「從 AI 出發的原生公司」。我們的日常營運本身就是 AI Agent 協作的"
    "最佳實踐，這意味著我們提供的解決方案，是經過自身驗證的成熟方案，而非紙上談兵。"
)

add_heading_styled("完整 Agent 協作框架", level=2)
add_para(
    "我們擁有完整的 AI Agent 開發、部署與管理框架，包含 Agent 設計模式庫、協作調度引擎、"
    "監控管理平台等，能快速且穩定地為客戶打造客製化的 AI Agent 解決方案。"
)

add_heading_styled("Palantir 級企業解決方案", level=2)
add_para(
    "凌策公司採用與 Palantir AIP 同等級的企業 AI 平台架構，具備處理大型企業複雜業務場景的能力，"
    "同時在資料安全與合規方面達到企業級標準。我們的方案不僅具備技術深度，更擁有企業級的穩定性與可靠性。"
)

why_table = [
    ["比較項目", "傳統 SI 廠商", "一般 AI 新創", "凌策公司"],
    ["AI 技術深度", "淺層應用", "單一技術", "全棧 AI Agent"],
    ["企業理解", "深", "淺", "AI 原生 + 企業理解"],
    ["實施速度", "3-6 個月", "1-3 個月", "4 週快速導入"],
    ["方案成熟度", "拼裝整合", "POC 階段", "自身驗證的成熟方案"],
    ["持續演進", "版本升級", "技術迭代", "AI Agent 自主學習優化"],
]
make_table(why_table[0], why_table[1:], col_widths=[3, 3.5, 3.5, 5])

doc.add_page_break()

# ============================================================
# 9. CONTACT INFORMATION
# ============================================================
add_heading_styled("八、聯絡資訊", level=1)

add_para(
    "感謝維明撥冗閱覽本提案書。凌策公司期待與維明攜手合作，共同開創 AI 驅動的企業服務新時代。"
    "如有任何疑問或進一步的合作意向，歡迎隨時與我們聯繫。"
)

doc.add_paragraph()

contact_data = [
    ["項目", "資訊"],
    ["公司名稱", "凌策公司（LingCe Company）"],
    ["聯絡窗口", "凌策商務 Agent"],
    ["電子郵件", "contact@lingce.com"],
    ["公司網站", "www.lingce.com"],
    ["服務時間", "24/7（AI Agent 全天候服務）"],
]
make_table(contact_data[0], contact_data[1:], col_widths=[4, 12])

doc.add_paragraph()
doc.add_paragraph()

# Closing
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("━" * 30)
set_run_font(run, size=12, color=BLUE)

add_para(
    "「以 AI 之力，成就企業之美」",
    size=14, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)
add_para(
    "凌策公司 LingCe Company",
    size=12, bold=False, color=DARK_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("本文件為凌策公司機密文件，未經授權不得複製或散佈。")
set_run_font(run, size=8, color=RGBColor(0x99, 0x99, 0x99))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Copyright 2026 LingCe Company. All Rights Reserved.")
set_run_font(run, size=8, color=RGBColor(0x99, 0x99, 0x99))

# ============================================================
# SAVE
# ============================================================
output_path = r"C:\Users\B00325\Desktop\公司AI比賽\lingce-company\docs\proposals\維明_提案書.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Proposal saved to: {output_path}")
